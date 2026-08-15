import logging
from pathlib import Path
import pdfplumber
import pypdf
import docx

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract raw text from a PDF file using pdfplumber with pypdf fallback."""
    text_chunks = []
    
    # Try pdfplumber first
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
        
        extracted = "\n\n".join(text_chunks).strip()
        if extracted:
            return extracted
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed on {file_path.name}: {e}. Falling back to pypdf.")
    
    # Fallback to pypdf
    try:
        reader = pypdf.PdfReader(file_path)
        pypdf_chunks = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pypdf_chunks.append(page_text)
        return "\n\n".join(pypdf_chunks).strip()
    except Exception as e:
        logger.error(f"pypdf extraction failed on {file_path.name}: {e}")
        raise ValueError(f"Could not extract text from PDF: {e}")

def extract_text_from_docx(file_path: Path) -> str:
    """Extract raw text from a .docx file including paragraphs and tables."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
                
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed on {file_path.name}: {e}")
        raise ValueError(f"Could not extract text from DOCX document: {e}")

def extract_cv_text(file_path: Path) -> str:
    """Extract text from supported CV file formats (.pdf, .docx, .doc)."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")
